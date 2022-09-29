"""Tests for common.document_parser.lib.section_parse.utils.utils"""

from unittest import TestCase, main
from os.path import dirname, join
from docx import Document
from collections import defaultdict
from copy import deepcopy
from typing import List, Callable
import sys

sys.path.append(dirname(__file__).replace("/section_parse/tests/unit", ""))
from section_parse.tests import TestItem
from section_parse.utils import (
    match_section_num,
    match_roman_numerals,
    match_enclosure_num,
    match_num_list_item,
    match_num_dot,
    match_num_parentheses,
    is_alpha_list_item,
    is_sentence_continuation,
    is_toc,
    is_glossary_continuation,
    is_next_num_list_item,
    is_list_child,
    is_bold,
    is_first_line_indented,
    remove_strikethrough_text,
    starts_with_bullet,
    match_attachment_num,
    is_attachment_start,
    is_subsection_start_for_section_1,
    get_subsection_of_section_1,
)


class UtilsTest(TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        test_doc_path = join(
            dirname(__file__).replace("unit", "data"), "test_utils.docx"
        )
        cls.test_doc = Document(test_doc_path)

        # Set up for test_remove_strikethrough_text().
        paragraphs = list(cls.test_doc.paragraphs)
        # Keys are paragraph index, values are list of text with strikethrough
        # in that paragraph.
        cls.paragraphs_with_strike = defaultdict(list)
        for i in range(len(paragraphs)):
            paragraph = paragraphs[i]
            for run in paragraph.runs:
                if run.font.strike:
                    cls.paragraphs_with_strike[i].append(run.text)

        if len(cls.paragraphs_with_strike) == 0:
            raise Exception(
                f"ERROR: Cannot run `test_remove_strikethrough_text()` because "
                f"no strikethrough text was detected in the test file: "
                f"`{test_doc_path}`."
            )

    def _run(self, func: Callable, test_cases: List[TestItem]):
        for test in test_cases:
            test.set_func(func)
            test.set_test_obj(self)
            test.verify_output()

    def test_match_section_num(self):
        """Verifies match_section_num()"""
        test_cases = [
            TestItem(("Section 5", None), "5"),
            TestItem(("Section 5", "5"), "5"),
            TestItem(("hello section 4", "12"), None),
            TestItem(("hello section 7", None), None),
            TestItem(("Section 2", "4"), None),
            TestItem(("A.2. Procedures", None), "A.2"),
            TestItem(("1.14", None), "1"),
            TestItem(("1.14", "2.14"), None),
        ]
        self._run(match_section_num, test_cases)

    def test_match_roman_numerals(self):
        """Verifies match_roman_numerals()."""
        test_cases = [
            TestItem(["III."], 3),
            TestItem(["II."], 2),
            TestItem(["XLARGE"], None),
        ]
        self._run(match_roman_numerals, test_cases)

    def test_match_enclosure_num(self):
        """Verifies match_enclosure_num()."""
        with self.assertRaises(ValueError):
            match_enclosure_num("", return_type="hi")
            match_enclosure_num("", return_type="string")
            match_enclosure_num("", return_type="int")
            match_enclosure_num("", return_type="list")

        test_cases = [
            TestItem(("Enclosure 4", None, "str"), "4"),
            TestItem(("Enclosure 3", "4", "bool"), False),
            TestItem(("Enclosure 4", "4", "bool"), True),
            TestItem(("E1.  Enclosure 1", None, "str"), "1"),
        ]
        self._run(match_enclosure_num, test_cases)

    def test_match_num_list_item(self):
        """Verifies match_num_list_item()."""
        test_cases = [
            TestItem(["1."], "1"),
            TestItem(["(2)"], "2"),
            TestItem(["3)"], "3"),
            TestItem(["hello 1."], None),
        ]
        self._run(match_num_list_item, test_cases)

    def test_match_num_dot(self):
        """Verifies match_num_dot()."""
        test_cases = [
            TestItem(("12.",), "12"),
            TestItem(("3",), None),
            TestItem(("hello 4.",), None),
        ]
        self._run(match_num_dot, test_cases)

    def test_match_num_parentheses(self):
        """Verifies match_num_parentheses()."""
        test_cases = [
            TestItem(("(1)",), "1"),
            TestItem(("2)",), "2"),
            TestItem(("3",), None),
            TestItem(("hello (1)",), None),
        ]
        self._run(match_num_parentheses, test_cases)

    def test_is_bold(self):
        """Verifies is_bold()."""
        expected_bold_indices = [4, 7, 10, 13, 20]
        test_cases = []

        for i in range(len(self.test_doc.paragraphs)):
            par = self.test_doc.paragraphs[i]
            test_cases.append(TestItem((par,), i in expected_bold_indices))

        self._run(is_bold, test_cases)

    def test_is_first_line_indented(self):
        """Verifies is_first_line_indented()."""
        expected_indent_indices = [5, 8, 11]
        test_cases = []
        for i in range(len(self.test_doc.paragraphs)):
            par = self.test_doc.paragraphs[i]
            test_cases.append(TestItem((par,), i in expected_indent_indices))

        self._run(is_first_line_indented, test_cases)

    def test_is_sentence_continuation(self):
        """Verifies is_sentence_continuation()."""
        test_cases = [
            TestItem(("hello", "she says "), True),
            TestItem(("hello", "she says."), False),
            TestItem(("Hello", "she says "), False),
        ]
        self._run(is_sentence_continuation, test_cases)

    def test_is_alpha_list_item(self):
        """Verifies is_alpha_list_item()."""
        test_cases = [
            TestItem(("a. Responsibilities",), True),
            TestItem(("aa. Procedures",), True),
            TestItem(("(a) Policy",), True),
            TestItem(("a Reference",), False),
            TestItem(("reference (a)",), False),
        ]
        self._run(is_alpha_list_item, test_cases)

    def test_is_toc(self):
        """Verifies is_toc()."""
        test_cases = [
            TestItem(("PROCEDURES...........5",), True),
            TestItem(("DoD 5 TABLE OF CONTENTS",), True),
            TestItem((" Table of Contents",), True),
        ]
        self._run(is_toc, test_cases)

    def test_is_glossary_continuation(self):
        """Verifies is_glossary_continuation()."""
        test_cases = [
            TestItem(("5\tGLOSSARY ",), True),
            TestItem(("G.4 ",), True),
            TestItem(("50 GLOSSARY G.5. DoDI 1000.4",), True),
        ]
        self._run(is_glossary_continuation, test_cases)

    def test_is_next_num_list_item(self):
        """Verifies is_next_num_list_item()."""
        test_cases = [
            TestItem(
                (
                    "3.  Complainant.",
                    ["1.  Informal Complaint. ", "2.  Formal Complaint."],
                ),
                True,
            ),
            TestItem(
                (
                    "1.  DoD Military Equal Opportunity (MEO) ",
                    ["6.  Equal Opportunity (EO)"],
                ),
                False,
            ),
            TestItem(
                (
                    "(4)  The provision of health",
                    [
                        "(2)  Non-DoD government agencies ",
                        "(3)  The Armed Forces Repository",
                    ],
                ),
                True,
            ),
            TestItem(
                (
                    "4. Subsection B of",
                    [
                        "3.  blah.",
                        "(1)  OSD, the Military Departments ",
                        "(2)  Non-DoD government agencies",
                    ],
                ),
                True,
            ),
        ]
        self._run(is_next_num_list_item, test_cases)

    def test_is_list_child(self):
        """Verifies is_list_child()."""
        test_cases = [
            TestItem(
                (
                    "III.  Point(s) of Contact",
                    [
                        "I.  Background and Current Situation",
                        "II.  Operations Status and Update:",
                    ],
                ),
                True,
            ),
            TestItem(("1.2.  POLICY.", ["1.1.  APPLICABILITY"]), True),
            TestItem(
                ("E2.1.2.  Complaint.", ["E2.1.1.  Affirmative Action. "]),
                True,
            ),
        ]
        self._run(is_list_child, test_cases)

    def test_remove_strikethrough_text(self):
        """Verifies remove_strikethrough_text()."""
        fail_msg = "FAILURE: test_remove_strikethrough_text()."

        paragraphs = deepcopy(list(self.test_doc.paragraphs))
        for i in range(len(paragraphs)):
            text_before = paragraphs[i].text
            remove_strikethrough_text(paragraphs[i])
            text_after = paragraphs[i].text

            if i in self.paragraphs_with_strike.keys():
                for strike_text in self.paragraphs_with_strike[i]:
                    self.assertTrue(
                        strike_text not in text_after,
                        fail_msg + f"Strikethrough text was detected in the "
                        f"paragraph after calling `remove_strikethrough_text().`",
                    )
            else:
                self.assertEqual(
                    text_before,
                    text_after,
                    fail_msg + f"No strikethrough was detected in the text, "
                    f"but the text was somehow altered. Text before was: "
                    f"`{text_before}`. Text after was: `{text_after}`.",
                )

    def test_starts_with_bullet(self):
        """Verifies starts_with_bullet()."""
        test_cases = [
            TestItem(("•First",), True),
            TestItem(("oNext",), True),
        ]
        self._run(starts_with_bullet, test_cases)

    def test_match_attachment_num(self):
        """Verifies match_attachment_num()."""
        test_cases = [
            TestItem(("ATTACHMENT 1 RESPONSIBILITIES",), "1"),
            TestItem(("ATTACHMENT REFERENCES",), None),
            TestItem(("Attachment 22",), "22"),
            TestItem(("This attachment 12",), None),
        ]
        self._run(match_attachment_num, test_cases)

    def test_is_attachment_start(self):
        """Verifies is_attachment_start()."""
        test_cases = [
            TestItem(("ATTACHMENT ",), True),
            TestItem(("See Attachment 12",), False),
            TestItem(("ATTACHMENT 4 PROCEDURES",), True),
        ]
        self._run(is_attachment_start, test_cases)

    def test_is_subsection_start_for_section_1(self):
        """Verifies is_subsection_start_for_section_1()."""
        test_cases = [
            TestItem(("1.1   Reissues reference", ""), True),
            TestItem(("APPLICABILITY. ", ""), True),
            TestItem(("POLICY", ""), True),
            TestItem(("Information Collections", ""), True),
            TestItem(("applicability", ""), False),
            TestItem(("1.14. Award Resources", ""), True),
            TestItem(("1.2. POLICY", "policy"), True),
            TestItem(("1.6 Awards", "Awards"), True),
            TestItem(("BREAKING NEWS", ""), False),
            TestItem(("1.2 RESOURCES", "Policy"), False),
        ]

        self._run(is_subsection_start_for_section_1, test_cases)

    def test_get_subsection_of_section_1(self):
        """Verifies get_subsection_of_section_1()."""
        section_1 = [
            "SECTION 1:GENERAL INFORMATION",
            "1.1.  APPLICABILITY.  This issuance applies to OSD, the Military Departments, the Office of the Chairman of the Joint Chiefs of Staff and the Joint Staff, the Combatant Commands, the Office of the Inspector General of the Department of Defense (IG DoD), the Defense Agencies, the DoD Field Activities, and all other organizational entities within the DoD (referred to collectively in this issuance as the “DoD Components”). ",
            "1.2.  POLICY.",
            "a.  In accordance with Volume 451 of DoD Instruction 1400.25: ",
            "(1)  Secretary of Defense Honorary Awards are granted consistent with equal employment opportunity and policies, laws, regulations, and Executive orders that prohibit unlawful discrimination based on race, color, religion, sex, national origin, age, disability, genetic information, reprisal for protected activity, marital status, political affiliation, or any other unlawful factor. ",
            "(2)  Requests by non-DoD personnel either to nominate themselves or others or to endorse nominations for themselves or others for awards or decorations sponsored by the DoD, other federal agencies, or private organizations, will not be honored. ",
            "b.  Secretary of Defense Honorary Award nominations must be endorsed as follows: ",
            "\t(1)  The OSD Principal Staff Assistants must endorse award nominations originating within their respective OSD Components. ",
            "\t(2)  The Chairman of the Joint Chiefs of Staff must endorse award nominations originating within a Combatant Command and the Joint Staff. ",
            "(3)  Award nominations originating within a Defense Agency or DoD Field Activity must be endorsed by the OSD Principal Staff Assistant who has authority, direction, and control over the Defense Agency, DoD Field Activity, or other organizational entity concerned. ",
            "\t(4)  Secretaries of the Military Departments must endorse award nominations originating within their respective departments. ",
            "\t(5)  The Chief of the National Guard Bureau must endorse award nominations originating within the National Guard Bureau. ",
            "1.3.  INFORMATION COLLECTIONS.  Secretary of Defense Honorary Awards for DoD civilian employees referred to in this issuance do not require licensing with a report control symbol in accordance with Paragraph 2.b.(2) of Volume 1 of DoD Manual 8910.01. ",
            "DoDM 1432.04, August 10, 2018",
            "1.4.  AWARD RESOURCES.  ",
            "Nomination templates for awards covered under this issuance may be obtained by contacting the Washington Headquarters Services (WHS), Human Resources Directorate (HRD), Performance Management and Awards Division (PM&AD). ",
        ]
        test_cases = [
            TestItem((section_1, "Applicability"), [section_1[1]]),
            TestItem((section_1, "Policy"), section_1[2:12]),
            TestItem((section_1, "Information Collections"), section_1[12:14]),
            TestItem((section_1, "award resources"), section_1[14:]),
        ]
        self._run(get_subsection_of_section_1, test_cases)


if __name__ == "__main__":
    main(failfast=True)
